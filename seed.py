import os
import logging
import json
from pathlib import Path
from langchain_community.vectorstores import Weaviate
from langchain_community.embeddings import HuggingFaceInferenceAPIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter
from mcp_host.config import Settings
import weaviate
import re
from typing import List, Dict, Any, Tuple
from langchain.schema import Document
import uuid
from docx import Document as DocxDocument

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Simple sentence splitter (no nltk dependency)
def simple_sent_tokenize(text: str) -> List[str]:
    """Split text into sentences using regex."""
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    return [s.strip() for s in sentences if s.strip()]

# Optional: Try to load spaCy for entity extraction
try:
    import spacy
    nlp = spacy.load('en_core_web_sm')
    SPACY_AVAILABLE = True
except (ImportError, OSError):
    logging.warning("spaCy not available. Entity extraction disabled.")
    nlp = None
    SPACY_AVAILABLE = False

def extract_entities(text: str) -> Dict[str, List[str]]:
    """
    Extracts named entities from text using spaCy.
    Returns empty dict if spaCy model is not available.
    """
    if nlp is None:
        return {}
    
    doc = nlp(text)
    entities = {}
    for ent in doc.ents:
        if ent.label_ not in entities:
            entities[ent.label_] = []
        entities[ent.label_].append(ent.text)
    
    # Deduplicate
    for label in entities:
        entities[label] = sorted(list(set(entities[label])))
        
    return entities

def load_documents_from_directory(directory_path: str) -> List[Document]:
    """
    Load documents from directory using direct format-specific loaders.
    Supports: .docx, .txt, .md
    No unstructured dependency - robust and fast.
    """
    documents = []
    path = Path(directory_path)
    
    # Supported extensions
    supported_files = list(path.rglob('*.docx')) + list(path.rglob('*.txt')) + list(path.rglob('*.md'))
    
    logging.info(f"Found {len(supported_files)} supported files")
    
    for file_path in supported_files:
        try:
            if file_path.suffix == '.docx':
                # Load DOCX with heading-aware section detection
                full_text, sections = _detect_sections_docx(file_path)
                documents.append(Document(
                    page_content=full_text,
                    metadata={"source": str(file_path), "type": "docx", "sections": sections}
                ))
                logging.info(f"✓ Loaded {file_path.name} ({len(full_text)} chars, {len(sections)} sections)")
                
            elif file_path.suffix in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                sections = _detect_sections_text(text, file_path.suffix)
                documents.append(Document(
                    page_content=text,
                    metadata={"source": str(file_path), "type": file_path.suffix[1:], "sections": sections}
                ))
                logging.info(f"✓ Loaded {file_path.name} ({len(text)} chars, {len(sections)} sections)")
                
        except Exception as e:
            logging.error(f"✗ Failed to load {file_path.name}: {e}")
            continue
    
    return documents


# ============================================================================
# SECTION DETECTION — 2-stage LangChain pipeline
#   Stage 1: MarkdownHeaderTextSplitter  → structure-aware section splitting
#   Stage 2: RecursiveCharacterTextSplitter → sub-chunk large sections
# ============================================================================

# LangChain Stage 1 splitter — splits by heading hierarchy.
# .docx and .txt are first converted to markdown format, then split here.
_md_header_splitter = MarkdownHeaderTextSplitter(
    headers_to_split_on=[
        ("#", "Header 1"),
        ("##", "Header 2"),
        ("###", "Header 3"),
        ("####", "Header 4"),
    ],
)


def _is_heading_by_heuristic(para) -> bool:
    """
    Auto-detect whether a python-docx paragraph is a heading using
    multiple heuristics — works even if the document has no formal
    Heading styles applied.

    Checks (any one is enough):
      1. Word style name starts with "Heading" or equals "Title" / "Subtitle"
      2. Entire paragraph is bold and ≤ 120 chars
      3. Font size ≥ 14 pt and ≤ 120 chars
      4. Text is ALL CAPS, has ≥ 2 words, and ≤ 120 chars
      5. Text matches common numbered-heading patterns
         ("1.", "1.1", "A.", "I.", "Chapter 3", "Section 2", etc.)
    """
    text = para.text.strip()
    if not text or len(text) > 120:
        return False

    # 1. Formal Word style
    style = (para.style.name or "") if para.style else ""
    if style.startswith("Heading") or style in ("Title", "Subtitle"):
        return True

    word_count = len(text.split())
    if word_count < 1 or word_count > 15:
        return False            # Too long for a heading

    # 2. All runs bold
    runs = [r for r in para.runs if r.text.strip()]
    if runs and all(r.bold for r in runs):
        return True

    # 3. Large font (≥ 14pt)
    sizes = {r.font.size.pt for r in runs if r.font and r.font.size}
    if sizes and min(sizes) >= 14:
        return True

    # 4. ALL CAPS with at least 2 words (rules out acronyms like "AI")
    if text == text.upper() and word_count >= 2 and any(c.isalpha() for c in text):
        return True

    # 5. Numbered heading patterns
    if re.match(
        r'^(?:'
        r'\d{1,3}(?:\.\d{1,3}){0,3}\.?\s+'
        r'|[A-Z]\.\s+'
        r'|[IVXLC]+\.\s+'
        r'|(?:Chapter|Section|Part|Appendix|Annex|Module|Unit|Phase|Pillar|Pillar\s*\d)'
        r')'
        , text, re.IGNORECASE
    ):
        return True

    return False


def _infer_heading_level(para) -> int:
    """Infer markdown heading level from a python-docx paragraph style."""
    style = (para.style.name or "") if para.style else ""
    if style == "Title":
        return 1
    if style == "Subtitle":
        return 2
    if style.startswith("Heading"):
        try:
            return min(int(style.replace("Heading", "").strip()), 4)
        except (ValueError, IndexError):
            return 1
    # Heuristic-detected headings (bold, ALL CAPS, etc.) → level 1
    return 1


def _docx_to_markdown(file_path) -> Tuple[str, str]:
    """
    Convert a .docx to markdown-formatted text so LangChain's
    MarkdownHeaderTextSplitter can do structure-aware splitting.

    Returns (markdown_text, plain_full_text).
    """
    doc = DocxDocument(str(file_path))
    md_lines: List[str] = []
    plain_lines: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        plain_lines.append(text)

        if _is_heading_by_heuristic(para):
            level = _infer_heading_level(para)
            md_lines.append(f"\n{'#' * level} {text}\n")
        else:
            md_lines.append(text)

    return '\n'.join(md_lines), '\n'.join(plain_lines)


def _detect_sections_docx(file_path) -> Tuple[str, List[Dict[str, str]]]:
    """
    Parse .docx → convert to markdown → split with LangChain's
    MarkdownHeaderTextSplitter for automatic structure-aware sections.

    Two LangChain splitters work in a pipeline:
      Stage 1: MarkdownHeaderTextSplitter  → sections by heading
      Stage 2: RecursiveCharacterTextSplitter → sub-chunks (in _build_sub_chunks)
    """
    md_text, full_text = _docx_to_markdown(file_path)
    section_docs = _md_header_splitter.split_text(md_text)

    sections: List[Dict[str, str]] = []
    for doc in section_docs:
        # Pick the most specific (deepest) heading level
        heading = (doc.metadata.get("Header 4") or doc.metadata.get("Header 3")
                   or doc.metadata.get("Header 2") or doc.metadata.get("Header 1")
                   or Path(file_path).stem)
        body = doc.page_content.strip()
        if body:
            sections.append({"heading": heading, "text": body})

    # Fallback: entire document as a single section
    if not sections:
        sections.append({"heading": Path(file_path).stem, "text": full_text})

    logging.info(f"  → Detected {len(sections)} sections in {Path(file_path).name}: "
                 f"{[s['heading'][:50] for s in sections]}")
    return full_text, sections


def _is_text_heading(line: str) -> bool:
    """
    Auto-detect heading lines in plain text / markdown.
    Matches: # headings, ALL CAPS lines, numbered headings, underline patterns.
    """
    stripped = line.strip()
    if not stripped or len(stripped) > 120:
        return False

    # Markdown # heading
    if re.match(r'^#{1,4}\s+', stripped):
        return True

    word_count = len(stripped.split())
    if word_count < 1 or word_count > 15:
        return False

    # ALL CAPS (at least 2 words)
    if stripped == stripped.upper() and word_count >= 2 and any(c.isalpha() for c in stripped):
        return True

    # Numbered heading
    if re.match(
        r'^(?:'
        r'\d{1,3}(?:\.\d{1,3}){0,3}\.?\s+'
        r'|[A-Z]\.\s+'
        r'|[IVXLC]+\.\s+'
        r'|(?:Chapter|Section|Part|Appendix|Annex|Module|Unit|Phase)'
        r')'
        , stripped, re.IGNORECASE
    ):
        return True

    return False


def _detect_sections_text(text: str, file_suffix: str = ".md") -> List[Dict[str, str]]:
    """
    Split .md / .txt into sections using LangChain's MarkdownHeaderTextSplitter.

    - .md files  → fed directly to the splitter (already has # headings).
    - .txt files → auto-detected headings converted to # format first,
                    then split by LangChain.
    """
    if file_suffix == ".md":
        md_text = text
    else:
        # Convert plain-text headings to markdown # format
        lines = text.split('\n')
        md_lines: List[str] = []
        i = 0
        while i < len(lines):
            stripped = lines[i].strip()

            # Underline-style heading: "Title\n=====" → "# Title"
            if stripped and i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if next_line and len(next_line) >= 3:
                    if set(next_line) <= {'='}:
                        md_lines.append(f"\n# {stripped}\n")
                        i += 2
                        continue
                    elif set(next_line) <= {'-'}:
                        md_lines.append(f"\n## {stripped}\n")
                        i += 2
                        continue

            # Other heading patterns (ALL CAPS, numbered, etc.)
            if stripped and _is_text_heading(stripped) and not stripped.startswith('#'):
                md_lines.append(f"\n# {stripped}\n")
            else:
                md_lines.append(lines[i])
            i += 1
        md_text = '\n'.join(md_lines)

    section_docs = _md_header_splitter.split_text(md_text)

    sections: List[Dict[str, str]] = []
    for doc in section_docs:
        heading = (doc.metadata.get("Header 4") or doc.metadata.get("Header 3")
                   or doc.metadata.get("Header 2") or doc.metadata.get("Header 1")
                   or "Content")
        body = doc.page_content.strip()
        if body:
            sections.append({"heading": heading, "text": body})

    if not sections:
        sections.append({"heading": "Content", "text": text})

    return sections


# ============================================================================
# EXTRACTIVE SUMMARIES
# ============================================================================

def generate_extractive_summary(
    text: str, max_sentences: int = 3, max_tokens: int = 150
) -> str:
    """Create an extractive summary from the first N sentences, capped at max_tokens."""
    sentences = simple_sent_tokenize(text)
    parts: List[str] = []
    token_count = 0
    for sent in sentences[:max_sentences]:
        sent_tokens = len(sent.split())
        if token_count + sent_tokens > max_tokens and parts:
            break
        parts.append(sent)
        token_count += sent_tokens
    return ' '.join(parts) if parts else text[:400]


# ============================================================================
# SEMANTIC HIERARCHICAL CHUNKING
# ============================================================================

# LangChain Stage 2 splitter — recursive sub-chunk splitting.
# Breaks large sections into overlapping sub-chunks for embedding.
_langchain_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1024,          # ~256 tokens × 4 chars/token
    chunk_overlap=256,        # ~64 tokens overlap
    separators=["\n\n", "\n", ". ", "? ", "! ", "; ", ", ", " ", ""],
    length_function=len,
    is_separator_regex=False,
)


def _build_sub_chunks(
    text: str,
    chunk_size: int = 1024,
    chunk_overlap: int = 256,
) -> List[str]:
    """
    Split a long section into sub-chunks using LangChain's
    RecursiveCharacterTextSplitter — respects sentence and paragraph
    boundaries automatically.
    """
    # Use the module-level splitter if defaults match, else create one
    if chunk_size == 1024 and chunk_overlap == 256:
        splitter = _langchain_splitter
    else:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", "? ", "! ", "; ", ", ", " ", ""],
            length_function=len,
            is_separator_regex=False,
        )
    return splitter.split_text(text)


def semantic_hierarchical_chunking(
    documents: List[Document],
    section_max_chars: int = 2048,
    sub_chunk_size: int = 1024,
    sub_chunk_overlap: int = 256,
) -> List[Document]:
    """
    Semantic hierarchical chunking with three levels:

    Level 0 — One document-summary chunk per file (extractive summary).
    Level 1 — One chunk per auto-detected section (kept whole if ≤ section_max_chars).
    Level 2 — Sub-chunks via LangChain RecursiveCharacterTextSplitter when section is large.

    Every chunk carries an extractive summary and parent_id so the retriever
    can walk the hierarchy (child → section → document).

    Works with ANY document — headings are auto-detected from formatting,
    bold, font size, ALL CAPS, or numbered patterns.
    """
    logging.info(
        f"Starting semantic hierarchical chunking: "
        f"section_max={section_max_chars} chars, sub_chunk={sub_chunk_size} chars, overlap={sub_chunk_overlap} chars"
    )
    all_chunks: List[Document] = []

    for doc in documents:
        full_text = doc.page_content
        source = doc.metadata.get("source", "unknown")
        sections = doc.metadata.get("sections", [])

        # ── Level 0: Document summary ───────────────────────────────────────
        doc_summary = generate_extractive_summary(full_text, max_sentences=5, max_tokens=200)
        doc_chunk_id = str(uuid.uuid4())
        all_chunks.append(Document(
            page_content=doc_summary,
            metadata={
                "chunk_id": doc_chunk_id,
                "source": source,
                "level": 0,
                "section_title": "Document Summary",
                "summary": doc_summary,
                "parent_id": "",
                "chunk_index": 0,
                "total_chunks": 1,
                "chunk_size": len(doc_summary),
                "token_count": len(doc_summary.split()),
            }
        ))

        # Fall back to whole-document-as-one-section if no headings detected
        if not sections:
            sections = [{"heading": "Content", "text": full_text}]

        # ── Level 1 + Level 2 ──────────────────────────────────────────────
        for sec_idx, section in enumerate(sections):
            heading = section["heading"]
            sec_text = section["text"].strip()
            if not sec_text:
                continue

            sec_tokens = len(sec_text.split())
            sec_chars  = len(sec_text)
            sec_summary = generate_extractive_summary(sec_text, max_sentences=3, max_tokens=120)
            section_chunk_id = str(uuid.uuid4())

            if sec_chars <= section_max_chars:
                # Section fits in one chunk → Level 1 (full text)
                all_chunks.append(Document(
                    page_content=sec_text,
                    metadata={
                        "chunk_id": section_chunk_id,
                        "source": source,
                        "level": 1,
                        "section_title": heading,
                        "summary": sec_summary,
                        "parent_id": doc_chunk_id,
                        "chunk_index": sec_idx,
                        "total_chunks": 1,
                        "chunk_size": len(sec_text),
                        "token_count": sec_tokens,
                    }
                ))
            else:
                # Section too large → Level 1 summary-only + Level 2 sub-chunks
                all_chunks.append(Document(
                    page_content=sec_summary,
                    metadata={
                        "chunk_id": section_chunk_id,
                        "source": source,
                        "level": 1,
                        "section_title": heading,
                        "summary": sec_summary,
                        "parent_id": doc_chunk_id,
                        "chunk_index": sec_idx,
                        "total_chunks": 0,
                        "chunk_size": len(sec_summary),
                        "token_count": len(sec_summary.split()),
                    }
                ))

                # Sub-chunk the section into Level 2 pieces
                sub_chunks = _build_sub_chunks(sec_text, sub_chunk_size, sub_chunk_overlap)

                for sub_idx, sub_text in enumerate(sub_chunks):
                    sub_summary = generate_extractive_summary(sub_text, max_sentences=2, max_tokens=80)
                    all_chunks.append(Document(
                        page_content=sub_text,
                        metadata={
                            "chunk_id": str(uuid.uuid4()),
                            "source": source,
                            "level": 2,
                            "section_title": heading,
                            "summary": sub_summary,
                            "parent_id": section_chunk_id,
                            "chunk_index": sub_idx,
                            "total_chunks": len(sub_chunks),
                            "chunk_size": len(sub_text),
                            "token_count": len(sub_text.split()),
                        }
                    ))

    l0 = sum(1 for c in all_chunks if c.metadata['level'] == 0)
    l1 = sum(1 for c in all_chunks if c.metadata['level'] == 1)
    l2 = sum(1 for c in all_chunks if c.metadata['level'] == 2)
    logging.info(f"✓ Created {len(all_chunks)} hierarchical chunks (L0={l0}, L1={l1}, L2={l2})")
    return all_chunks


def seed_documents_from_local():
    """
    Loads documents from a local directory, processes them with hierarchical chunking,
    and seeds them into Weaviate with enriched metadata for better retrieval.
    
    FEATURES:
    - Hierarchical chunking: Base chunks with metadata
    - Enriched metadata: summaries, chunk indices, hierarchy levels
    - Optimized for retrieval and synthesis
    """
    try:
        # Load settings from environment variables
        settings = Settings()
        logging.info("Loaded configuration settings.")

        knowledge_base_path = settings.KNOWLEDGE_BASE_PATH
        if not os.path.isdir(knowledge_base_path):
            logging.error(f"Knowledge base path '{knowledge_base_path}' is not a valid directory.")
            return

        # 1. Load documents from local directory using direct loaders (no unstructured)
        logging.info(f"Loading documents from local directory: {knowledge_base_path}")
        documents = load_documents_from_directory(knowledge_base_path)
        
        if not documents:
            logging.warning("No documents found in the specified local directory.")
            return
        logging.info(f"✓ Successfully loaded {len(documents)} document(s)")

        # 2. Create hierarchical chunks with multi-level summaries
        logging.info("Creating semantic chunks with multi-level summaries...")
        
        # Initialize embeddings with dual fallback: InferenceClient -> requests
        logging.info(f"Initializing embeddings with model {settings.EMBEDDING_MODEL}...")
        
        class MultiFallbackEmbeddings:
            """Embeddings with retry logic and 1024-dim models for best quality"""
            # Use 1024-dim models for better semantic richness and multilingual support
            PRIMARY_MODEL = "BAAI/bge-m3"  # 1024 dims - best quality
            FALLBACK_MODELS = [
                "BAAI/bge-large-en-v1.5",  # 1024 dims - fallback
            ]
            MAX_RETRIES = 3
            
            def __init__(self, api_key, model_name):
                self.api_key = api_key
                self.model_name = model_name
                self.working_model = None
                
            def _try_requests(self, text, model, timeout=60):
                import requests
                import time
                url = f"https://router.huggingface.co/hf-inference/models/{model}/pipeline/feature-extraction"
                headers = {"Authorization": f"Bearer {self.api_key}"}
                response = requests.post(url, headers=headers, json={"inputs": text}, timeout=timeout)
                response.raise_for_status()
                return response.json()
                
            def embed_query(self, text):
                import time
                
                # If we found a working model, try it first with retries
                if self.working_model:
                    for attempt in range(self.MAX_RETRIES):
                        try:
                            result = self._try_requests(text, self.working_model)
                            if result and len(result) > 0:
                                return result
                        except Exception as e:
                            if attempt < self.MAX_RETRIES - 1:
                                wait_time = 2 ** attempt  # Exponential backoff: 1s, 2s, 4s
                                logging.warning(f"⚠️ Retry {attempt+1}/{self.MAX_RETRIES} for {self.working_model} in {wait_time}s...")
                                time.sleep(wait_time)
                            else:
                                self.working_model = None  # Reset and try primary
                
                # Try primary model with retries
                for attempt in range(self.MAX_RETRIES):
                    try:
                        result = self._try_requests(text, self.PRIMARY_MODEL)
                        if result and len(result) > 0:
                            if self.working_model != self.PRIMARY_MODEL:
                                logging.info(f"✓ Embedding model: {self.PRIMARY_MODEL} (1024-dim)")
                                self.working_model = self.PRIMARY_MODEL
                            return result
                    except Exception as e:
                        if attempt < self.MAX_RETRIES - 1:
                            wait_time = 2 ** attempt
                            logging.warning(f"⚠️ Retry {attempt+1}/{self.MAX_RETRIES} for {self.PRIMARY_MODEL} in {wait_time}s: {e}")
                            time.sleep(wait_time)
                        else:
                            logging.warning(f"⚠️ Primary model {self.PRIMARY_MODEL} failed after {self.MAX_RETRIES} retries")
                
                # Try fallback models (same 1024-dim) with retries
                for model in self.FALLBACK_MODELS:
                    for attempt in range(self.MAX_RETRIES):
                        try:
                            result = self._try_requests(text, model)
                            if result and len(result) > 0:
                                if self.working_model != model:
                                    logging.info(f"✓ Fallback embedding model: {model} (1024-dim)")
                                    self.working_model = model
                                return result
                        except Exception as e:
                            if attempt < self.MAX_RETRIES - 1:
                                wait_time = 2 ** attempt
                                logging.warning(f"⚠️ Retry {attempt+1}/{self.MAX_RETRIES} for {model} in {wait_time}s: {e}")
                                time.sleep(wait_time)
                            else:
                                logging.warning(f"⚠️ Fallback model {model} failed after {self.MAX_RETRIES} retries")
                                break
                
                raise ValueError(f"All embedding models failed after retries")
                    
            def embed_documents(self, texts):
                return [self.embed_query(t) for t in texts]
        
        embeddings = MultiFallbackEmbeddings(
            api_key=settings.HUGGINGFACE_API_KEY,
            model_name=settings.EMBEDDING_MODEL
        )
        
        # Apply semantic hierarchical chunking
        chunks = semantic_hierarchical_chunking(documents)
        
        # Enrich chunks with extracted entities
        logging.info("Enriching chunks with named entities...")
        enriched_chunks = []
        for chunk in chunks:
            entities = extract_entities(chunk.page_content)
            if entities:
                chunk.metadata['entities'] = entities
            enriched_chunks.append(chunk)
        
        logging.info(f"✓ Enriched {len(enriched_chunks)} chunks with entity metadata.")
        
        # Log a few examples of enriched chunks
        for i, chunk in enumerate(enriched_chunks[:3]):
            if 'entities' in chunk.metadata:
                logging.info(f"  - Chunk {i+1} entities: {chunk.metadata['entities']}")
        
        # 🔍 CHUNK VISIBILITY LOGGING
        logging.info("\n" + "="*80)
        logging.info("📋 SEEDING VERIFICATION - ALL CHUNKS CREATED:")
        logging.info("="*80)
        for i, chunk in enumerate(enriched_chunks):
            logging.info(f"\n[CHUNK {i}]")
            logging.info(f"  Level: {chunk.metadata.get('level', '?')}")
            logging.info(f"  Section: {chunk.metadata.get('section_title', 'N/A')}")
            logging.info(f"  Size: {len(chunk.page_content)} chars ({chunk.metadata.get('token_count', '?')} tokens)")
            logging.info(f"  Source: {chunk.metadata.get('source', 'unknown')}")
            logging.info(f"  Summary: {chunk.metadata.get('summary', '')[:120]}")
            logging.info(f"  Entities: {chunk.metadata.get('entities', [])}")
            logging.info(f"  Content: {chunk.page_content[:250]}...")
        logging.info("="*80 + "\n")

        # 3. Seed Weaviate
        logging.info(f"Connecting to Weaviate at {settings.WEAVIATE_HOST}:{settings.WEAVIATE_PORT}...")
        client = weaviate.connect_to_local(
            host=settings.WEAVIATE_HOST,
            port=settings.WEAVIATE_PORT,
            grpc_port=settings.WEAVIATE_GRPC_PORT,
        )
        
        # Clear the collection to ensure fresh data
        collection_name = "Chunk"
        if client.collections.exists(collection_name):
            logging.warning(f"Collection '{collection_name}' already exists. Deleting and recreating.")
            client.collections.delete(collection_name)

        # Create collection with proper schema for new Weaviate v4 API
        from weaviate.classes.config import Configure
        client.collections.create(
            name=collection_name,
            vectorizer_config=Configure.Vectorizer.none(),  # We'll provide vectors
            vector_index_config=Configure.VectorIndex.hnsw()
        )
        logging.info(f"✓ Created Weaviate collection '{collection_name}'")

        # Seed documents with embeddings
        logging.info(f"Seeding {len(enriched_chunks)} hierarchically-structured chunks into Weaviate...")
        collection = client.collections.get(collection_name)
        
        # Generate embeddings and seed documents with better error handling
        for i, chunk in enumerate(enriched_chunks):
            try:
                logging.info(f"Embedding chunk {i+1}/{len(enriched_chunks)}...")
                vector = embeddings.embed_query(chunk.page_content)
                
                if not vector or len(vector) == 0:
                    logging.error(f"✗ Empty embedding returned for chunk {i}")
                    continue
                    
                collection.data.insert(
                    properties={
                        "content": chunk.page_content,
                        "source": chunk.metadata.get("source", ""),
                        "entities": json.dumps(chunk.metadata.get("entities", {})),
                        "summary": chunk.metadata.get("summary", ""),
                        "level": chunk.metadata.get("level", 1),
                        "section_title": chunk.metadata.get("section_title", ""),
                        "parent_id": chunk.metadata.get("parent_id", ""),
                    },
                    vector=vector
                )
                logging.info(f"✓ Seeded chunk {i+1} successfully")
            except Exception as e:
                logging.error(f"✗ Failed to seed chunk {i}: {e}")
                continue
        
        logging.info(f"✓ Successfully seeded Weaviate with {len(enriched_chunks)} document chunks (hierarchical structure with summaries).")

    except Exception as e:
        logging.error(f"An error occurred during the seeding process: {e}", exc_info=True)

if __name__ == "__main__":
    seed_documents_from_local()
